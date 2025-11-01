from __future__ import annotations
from io import BytesIO
from typing import Iterable, Sequence
from pathlib import Path

try:
    # python-docx (dependency of docxtpl) should already be present
    from docx import Document  # type: ignore
    from docx.shared import Pt, Cm, RGBColor  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except Exception as e:  # pragma: no cover - defensive import
    Document = None  # type: ignore

# Absolute path to the header image requested by the user
_HEADER_IMAGE = Path(r"D:\sh_school_015\assets\img\moraslat.png")


def _ensure_arabic_styles(doc):
    """Set basic RTL direction and Arabic-friendly font defaults.
    - Force Normal style to use Arial 11pt.
    - Set document default to RTL via style element (w:bidi) when possible.
    """
    try:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        # set east asia font for arabic compatibility
        rPr = style.element.rPr
        rFonts = rPr.rFonts
        rFonts.set(qn("w:eastAsia"), "Arial")
        font.size = Pt(11)
        # Set paragraph direction RTL on the style (w:bidi=1)
        pPr = style.element.get_or_add_pPr()
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            from docx.oxml import OxmlElement  # type: ignore
            bidi = OxmlElement(qn("w:bidi"))
            pPr.append(bidi)
        bidi.set(qn("w:val"), "1")
    except Exception:
        pass


def _apply_header_image(doc) -> None:
    """Insert header image.
    Modes:
      - full: fill the whole first page header area with the image (visually full-page header),
               keep compact header on subsequent pages to avoid pushing content on every page.
      - compact: small 2.5cm header image (previous behavior).
    Select via env ABS_DOCX_HEADER_MODE (full|compact), default 'full'.

    Implementation note: In 'full' mode we also attempt to place the header image
    'behind text' so that body content can appear on the same page when possible.
    """
    import os
    try:
      # Short-circuit: default to no header image unless explicitly enabled via env
      mode = os.environ.get("ABS_DOCX_HEADER_MODE", "none").lower().strip() or "none"
      if mode == "none" or not _HEADER_IMAGE.exists():
          return
      # Use the first section
      section = doc.sections[0]
      # Ensure we can customize first page header separately
      try:
          section.different_first_page_header_footer = True
      except Exception:
          pass

      # Helper to clear a header
      def _clear_header(hdr):
          try:
              hdr_elm = hdr._element  # noqa: SLF001
              for child in list(hdr_elm):
                  hdr_elm.remove(child)
          except Exception:
              try:
                  for p in hdr.paragraphs:
                      for r in list(p.runs):
                          r.clear()
                      p.text = ""
              except Exception:
                  pass

      # Helper to convert last added picture to "behind text" anchor
      def _try_place_behind_text(paragraph):
          try:
              from docx.oxml import OxmlElement  # type: ignore
              from docx.oxml.ns import nsmap  # type: ignore
              # Find the inline drawing element just added
              inline_elems = paragraph._p.xpath('.//w:drawing/wp:inline', namespaces=nsmap)  # type: ignore
              if not inline_elems:
                  return
              inline = inline_elems[-1]
              # Build an anchor element with behindDoc="1"
              anchor = OxmlElement('wp:anchor')
              # Copy size from inline extent if present
              extent = inline.find('wp:extent', namespaces=nsmap)
              if extent is not None:
                  new_extent = OxmlElement('wp:extent')
                  new_extent.set('cx', extent.get('cx'))
                  new_extent.set('cy', extent.get('cy'))
                  anchor.append(new_extent)
              anchor.set('behindDoc', '1')
              anchor.set('locked', '0')
              anchor.set('layoutInCell', '1')
              anchor.set('allowOverlap', '1')
              simplePos = OxmlElement('wp:simplePos')
              simplePos.set('x', '0'); simplePos.set('y', '0')
              anchor.append(simplePos)
              posH = OxmlElement('wp:positionH'); posH.set('relativeFrom', 'page')
              posH.append(OxmlElement('wp:align')).text = 'center'
              posV = OxmlElement('wp:positionV'); posV.set('relativeFrom', 'page')
              posV.append(OxmlElement('wp:align')).text = 'top'
              anchor.append(posH); anchor.append(posV)
              # Ensure the anchored image does not affect layout flow
              wrapNone = OxmlElement('wp:wrapNone')
              anchor.append(wrapNone)
              # Move graphic from inline to anchor
              graphic = inline.find('a:graphic', namespaces=nsmap)
              if graphic is not None:
                  graphic = graphic.clone()
                  anchor.append(graphic)
              # Replace inline with anchor
              parent = inline.getparent()
              parent.replace(inline, anchor)
          except Exception:
              # Best-effort; ignore failures
              pass

      if mode == "compact":
          # Compact header on all pages
          try:
              section.top_margin = Cm(1)
              section.bottom_margin = Cm(1.5)
              section.left_margin = Cm(1.5)
              section.right_margin = Cm(1.5)
              section.header_distance = Cm(0.2)
          except Exception:
              pass
          header = section.header
          _clear_header(header)
          para = header.add_paragraph()
          run = para.add_run()
          try:
              run.add_picture(str(_HEADER_IMAGE), height=Cm(2.5))
          except Exception:
              run.add_picture(str(_HEADER_IMAGE))
          try:
              para.alignment = WD_ALIGN_PARAGRAPH.CENTER
          except Exception:
              pass
          return

      # FULL mode:
      # First page header: image scaled to fill printable area
      try:
          # Use minimal margins so the visual impression is full-page
          section.top_margin = Cm(0)
          section.left_margin = Cm(0)
          section.right_margin = Cm(0)
          # Keep a small bottom margin for footers/body safety
          section.bottom_margin = Cm(1.5)
          section.header_distance = Cm(0)
      except Exception:
          pass
      # Dimensions (EMU length returned by python-docx Length); we can pass them directly
      page_w = getattr(section, "page_width", None)
      page_h = getattr(section, "page_height", None)
      left = getattr(section, "left_margin", Cm(0))
      right = getattr(section, "right_margin", Cm(0))
      top = getattr(section, "top_margin", Cm(0))
      bottom = getattr(section, "bottom_margin", Cm(0))
      avail_w = page_w - left - right if page_w else None
      avail_h = page_h - top - bottom if page_h else None
      # First page header
      fheader = getattr(section, "first_page_header", None) or section.header
      _clear_header(fheader)
      p = fheader.add_paragraph()
      run = p.add_run()
      try:
          if avail_w:
              # Scale by width only to avoid expanding page height
              run.add_picture(str(_HEADER_IMAGE), width=avail_w)
          else:
              run.add_picture(str(_HEADER_IMAGE))
      except Exception:
          # Fallback to width-based scaling
          try:
              if avail_w:
                  run.add_picture(str(_HEADER_IMAGE), width=avail_w)
              else:
                  run.add_picture(str(_HEADER_IMAGE))
          except Exception:
              pass
      # Best-effort: place the image behind text so body can overlap
      _try_place_behind_text(p)
      try:
          p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      except Exception:
          pass
      # All pages: apply full header image as well (per request)
      try:
          header = section.header
          _clear_header(header)
          para2 = header.add_paragraph()
          run2 = para2.add_run()
          try:
              if avail_w:
                  run2.add_picture(str(_HEADER_IMAGE), width=avail_w)
              else:
                  run2.add_picture(str(_HEADER_IMAGE))
          except Exception:
              try:
                  if avail_w:
                      run2.add_picture(str(_HEADER_IMAGE), width=avail_w)
                  else:
                      run2.add_picture(str(_HEADER_IMAGE))
              except Exception:
                  pass
          try:
              para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
          except Exception:
              pass
          # Try to place behind text on subsequent pages as well
          _try_place_behind_text(para2)
      except Exception:
          pass
    except Exception:
      # Fail silently; header is optional
      pass


def render_table_docx(headers: Sequence[str], rows: Iterable[Sequence[str]], title: str | None = None) -> bytes:
    """Render a simple DOCX file containing an optional title and a table with headers and rows.

    Ensures full RTL layout (document, paragraphs, tables) suitable for Arabic and a professional table layout:
    - Default style set to RTL in _ensure_arabic_styles.
    - All paragraphs right-aligned with w:bidi=1.
    - Tables flagged with w:bidiVisual and cells' paragraphs set RTL.
    - Table stretched to page width with sensible column widths.
    - Header row styled and repeated across pages.
    - Cell margins and vertical alignment adjusted; text wraps inside cells.

    Args:
        headers: list of column headers (strings)
        rows: iterable of rows (list/tuple of strings with same length as headers)
        title: optional document title (added as first paragraph)
    Returns:
        bytes of the generated .docx file
    """
    if Document is None:
        raise RuntimeError("python-docx is not available")

    # Local helpers for RTL & styling
    from docx.oxml import OxmlElement  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # type: ignore

    def _p_set_rtl(p):
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Compact spacing
            if p.paragraph_format:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                try:
                    p.paragraph_format.line_spacing = 1.15
                except Exception:
                    pass
        except Exception:
            pass
        try:
            pPr = p._p.get_or_add_pPr()  # type: ignore[attr-defined]
            bidi = pPr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement(qn("w:bidi"))
                pPr.append(bidi)
            bidi.set(qn("w:val"), "1")
        except Exception:
            pass

    doc = Document()
    _ensure_arabic_styles(doc)
    _apply_header_image(doc)

    # Enforce 1 cm side margins as requested (keep current top/bottom)
    try:
        section = doc.sections[0]
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    except Exception:
        pass

    if title:
        p = doc.add_paragraph()
        run = p.add_run(str(title))
        try:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(14)
        except Exception:
            pass
        _p_set_rtl(p)

    # Create table
    # Reverse columns order so the first logical column appears on the right side in Word
    headers = list(headers)
    headers_rev = list(reversed(headers))

    table = doc.add_table(rows=1, cols=len(headers_rev))
    table.style = "Table Grid"
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Disable autofit so assigned widths take effect
        table.autofit = False
    except Exception:
        pass

    # Mark table as RTL visual & full width
    try:
        section = doc.sections[0]
        content_width = section.page_width - section.left_margin - section.right_margin
        tbl = table._tbl  # noqa: SLF001
        tblPr = getattr(tbl, 'tblPr', None)
        if tblPr is None:
            tblPr = OxmlElement(qn('w:tblPr'))
            tbl.insert(0, tblPr)
        bidiVisual = OxmlElement(qn('w:bidiVisual'))
        bidiVisual.set(qn('w:val'), '1')
        tblPr.append(bidiVisual)
        # Also set table direction RTL if supported
        tblDir = OxmlElement(qn('w:tblDir'))
        tblDir.set(qn('w:val'), 'rtl')
        tblPr.append(tblDir)
        # Set table width
        tblW = OxmlElement(qn('w:tblW'))
        tblW.set(qn('w:w'), str(content_width))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
    except Exception:
        content_width = None  # type: ignore

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers_rev):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            _p_set_rtl(p)
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        # Header styling: bold + shading
        try:
            for r in hdr_cells[i].paragraphs[0].runs:
                r.font.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(11)
                # White header text
                try:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                except Exception:
                    pass
        except Exception:
            pass
        try:
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()  # noqa: SLF001
            shd = OxmlElement(qn('w:shd'))
            # Maroon background
            shd.set(qn('w:fill'), '800000')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
        except Exception:
            pass
    # Repeat header on each page
    try:
        trPr = table.rows[0]._tr.get_or_add_trPr()  # noqa: SLF001
        tblHeader = OxmlElement(qn('w:tblHeader'))
        tblHeader.set(qn('w:val'), '1')
        trPr.append(tblHeader)
    except Exception:
        pass

    # Column width heuristics based on Arabic labels
    def _suggest_width_cm(label: str) -> float:
        l = (label or '').strip()
        if any(k in l for k in ["الرقم", "الرقم المدرسي", "sid"]):
            return 2.0
        if any(k in l for k in ["اسم", "الطالب", "name"]):
            return 5.5
        if any(k in l for k in ["الصف", "class"]):
            return 3.0
        if any(k in l for k in ["ولي", "الأمر", "Parent"]):
            return 4.0
        if any(k in l for k in ["هاتف", "جوال", "Phone"]):
            return 3.5
        if any(k in l for k in ["نشط", "يحتاج", "active", "needs"]):
            return 2.0
        return 3.0

    try:
        widths_cm = [_suggest_width_cm(h) for h in headers_rev]
        total_cm = sum(widths_cm) or 1.0
        if content_width:
            # Convert content_width (dxa) to cm: 1 inch = 1440 dxa, 1 inch = 2.54 cm
            total_dxa = content_width
            cm_per_dxa = 2.54 / 1440.0
            target_cm = total_dxa * cm_per_dxa
            scale = target_cm / total_cm if total_cm else 1.0
        else:
            scale = 1.0
        from docx.shared import Cm as _Cm  # type: ignore
        for i, col in enumerate(table.columns):
            w_cm = widths_cm[i] * scale
            try:
                col.width = _Cm(w_cm)
            except Exception:
                pass
    except Exception:
        pass

    # Body rows
    for row in rows:
        cells = table.add_row().cells
        row_vals = list(reversed(list(row)))
        for i, val in enumerate(row_vals):
            cells[i].text = "" if val is None else str(val)
            # Vertical alignment and margins
            try:
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tcPr = cells[i]._tc.get_or_add_tcPr()  # noqa: SLF001
                # Set cell margins (top/bottom/right/left)
                tcMar = OxmlElement(qn('w:tcMar'))
                for side in ('top','bottom','right','left'):
                    mar = OxmlElement(qn(f'w:{side}'))
                    mar.set(qn('w:w'), '80')  # 80 twips ≈ 0.14 cm
                    mar.set(qn('w:type'), 'dxa')
                    tcMar.append(mar)
                tcPr.append(tcMar)
            except Exception:
                pass
            for p in cells[i].paragraphs:
                _p_set_rtl(p)
                try:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(11)
                except Exception:
                    pass

    # As a safety net, enforce RTL on any other paragraphs
    try:
        for para in doc.paragraphs:
            _p_set_rtl(para)
    except Exception:
        pass

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()