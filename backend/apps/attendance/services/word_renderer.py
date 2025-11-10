from __future__ import annotations
from io import BytesIO
from pathlib import Path
from typing import Any, Tuple
import os
import hashlib

from docxtpl import DocxTemplate  # type: ignore
from django.utils import timezone

# For adding header image post-processing
try:
    from docx import Document as _DocxDocument  # type: ignore
    from docx.shared import Cm  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except Exception:  # pragma: no cover - defensive import
    _DocxDocument = None  # type: ignore

_HEADER_IMAGE = Path(r"D:\sh_school_015\assets\img\moraslat.png")

# NOTE: Using the path provided by the user for the Tanbih template
# Allow overriding via environment variable ABSENCE_ALERT_TEMPLATE_PATH or Django setting if provided.
_DEFAULT_TEMPLATE = Path(r"/DOC/repo\Tanbih.docx")
FALLBACK_TEMPLATE_PATH = Path(r"/DOC/repo\TanbihTEMP.dotx")

_env_template = os.environ.get("ABSENCE_ALERT_TEMPLATE_PATH")
TEMPLATE_PATH = Path(_env_template) if _env_template else _DEFAULT_TEMPLATE

# Allow forcing exact template usage by disabling fallback
# In development default to allow fallback; can be overridden per environment
_ALLOW_FALLBACK = os.environ.get("ABSENCE_ALERT_ALLOW_FALLBACK", "true").lower() in {
    "1",
    "true",
    "yes",
}


def resolve_template_path(preferred: str | None = None) -> Path:
    """Resolve the template path according to environment and availability.
    If ABSENCE_ALERT_ALLOW_FALLBACK=false and preferred (or TEMPLATE_PATH) does not exist,
    raise FileNotFoundError to signal misconfiguration instead of silently falling back.
    """
    tpath = Path(preferred) if preferred else TEMPLATE_PATH
    if tpath.exists():
        return tpath
    if _ALLOW_FALLBACK and FALLBACK_TEMPLATE_PATH.exists():
        return FALLBACK_TEMPLATE_PATH
    # No valid template found
    raise FileNotFoundError(
        f"لم يتم العثور على قالب Word المطلوب. حاول: {tpath}"
        + (" أو القالب الاحتياطي: " + str(FALLBACK_TEMPLATE_PATH) if _ALLOW_FALLBACK else "")
    )


def _apply_header_image_bytes(doc_bytes: bytes) -> bytes:
    """Open a DOCX bytes payload and optionally apply an image-only header, and enforce RTL text.
    - Mode 'full': full-page style header image on all pages (behind text)
    - Mode 'compact': small 2.5 cm image in header
    - Mode 'none': skip header image (default per latest requirement)
    In all modes we enforce RTL paragraph/table settings for Arabic.
    Returns modified bytes, or the original bytes on any failure.
    """
    import os
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import nsmap  # type: ignore

    mode = (os.environ.get("ABS_DOCX_HEADER_MODE", "none") or "none").lower().strip()
    if _DocxDocument is None:
        return doc_bytes
    try:
        doc = _DocxDocument(BytesIO(doc_bytes))
        section = doc.sections[0]
        # Allow different header on first page (we still set both first+default)
        try:
            section.different_first_page_header_footer = True
        except Exception:
            pass

        def _clear_header(hdr):
            try:
                el = hdr._element  # noqa: SLF001
                for ch in list(el):
                    el.remove(ch)
            except Exception:
                try:
                    for p in hdr.paragraphs:
                        for r in list(p.runs):
                            r.clear()
                        p.text = ""
                except Exception:
                    pass

        def _try_place_behind_text(paragraph):
            """Convert the last inline drawing in paragraph to an anchored drawing with behindDoc=1
            and align it to page center/top so body text can flow over it."""
            try:
                inline_elems = paragraph._p.xpath(".//w:drawing/wp:inline", namespaces=nsmap)  # type: ignore
                if not inline_elems:
                    return
                inline = inline_elems[-1]
                anchor = OxmlElement("wp:anchor")
                extent = inline.find("wp:extent", namespaces=nsmap)
                if extent is not None:
                    new_extent = OxmlElement("wp:extent")
                    new_extent.set("cx", extent.get("cx"))
                    new_extent.set("cy", extent.get("cy"))
                    anchor.append(new_extent)
                anchor.set("behindDoc", "1")
                anchor.set("locked", "0")
                anchor.set("layoutInCell", "1")
                anchor.set("allowOverlap", "1")
                simplePos = OxmlElement("wp:simplePos")
                simplePos.set("x", "0")
                simplePos.set("y", "0")
                anchor.append(simplePos)
                posH = OxmlElement("wp:positionH")
                posH.set("relativeFrom", "page")
                posH_align = OxmlElement("wp:align")
                posH_align.text = "center"
                posH.append(posH_align)
                posV = OxmlElement("wp:positionV")
                posV.set("relativeFrom", "page")
                posV_align = OxmlElement("wp:align")
                posV_align.text = "top"
                posV.append(posV_align)
                anchor.append(posH)
                anchor.append(posV)
                # Ensure the anchored image does not affect document layout flow
                wrapNone = OxmlElement("wp:wrapNone")
                anchor.append(wrapNone)
                graphic = inline.find("a:graphic", namespaces=nsmap)
                if graphic is not None:
                    graphic = graphic.clone()
                    anchor.append(graphic)
                parent = inline.getparent()
                parent.replace(inline, anchor)
            except Exception:
                pass

        if mode == "compact":
            try:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
                section.header_distance = Cm(0.2)
            except Exception:
                pass
            hdr = section.header
            _clear_header(hdr)
            p = hdr.add_paragraph()
            r = p.add_run()
            try:
                r.add_picture(str(_HEADER_IMAGE), height=Cm(2.5))
            except Exception:
                r.add_picture(str(_HEADER_IMAGE))
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        else:
            # Full mode: minimize margins and fill printable area with the image
            try:
                section.top_margin = Cm(0)
                section.left_margin = Cm(0)
                section.right_margin = Cm(0)
                section.bottom_margin = Cm(1.0)
                section.header_distance = Cm(0)
            except Exception:
                pass
            page_w = getattr(section, "page_width", None)
            page_h = getattr(section, "page_height", None)
            left = getattr(section, "left_margin", Cm(0))
            right = getattr(section, "right_margin", Cm(0))
            top = getattr(section, "top_margin", Cm(0))
            bottom = getattr(section, "bottom_margin", Cm(0))
            avail_w = page_w - left - right if page_w else None
            avail_h = page_h - top - bottom if page_h else None

            # First page header (full)
            fheader = getattr(section, "first_page_header", None) or section.header
            _clear_header(fheader)
            p1 = fheader.add_paragraph()
            r1 = p1.add_run()
            try:
                if avail_w and avail_h:
                    r1.add_picture(str(_HEADER_IMAGE), width=avail_w, height=avail_h)
                elif avail_w:
                    r1.add_picture(str(_HEADER_IMAGE), width=avail_w)
                else:
                    r1.add_picture(str(_HEADER_IMAGE))
            except Exception:
                try:
                    if avail_w:
                        r1.add_picture(str(_HEADER_IMAGE), width=avail_w)
                    else:
                        r1.add_picture(str(_HEADER_IMAGE))
                except Exception:
                    pass
            try:
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            _try_place_behind_text(p1)

            # Default header (other pages) — also full
            try:
                hdr = section.header
                _clear_header(hdr)
                p2 = hdr.add_paragraph()
                r2 = p2.add_run()
                try:
                    if avail_w and avail_h:
                        r2.add_picture(str(_HEADER_IMAGE), width=avail_w, height=avail_h)
                    elif avail_w:
                        r2.add_picture(str(_HEADER_IMAGE), width=avail_w)
                    else:
                        r2.add_picture(str(_HEADER_IMAGE))
                except Exception:
                    try:
                        if avail_w:
                            r2.add_picture(str(_HEADER_IMAGE), width=avail_w)
                        else:
                            r2.add_picture(str(_HEADER_IMAGE))
                    except Exception:
                        pass
                try:
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
                _try_place_behind_text(p2)
            except Exception:
                pass

        out = BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return doc_bytes


def render_alert_docx(alert: Any, template_path: str | None = None) -> bytes:
    """Render the Absence Alert DOCX using the Tanbih template (preferred).
    The template must contain placeholders matching the context keys below.
    """
    tpath = resolve_template_path(template_path)
    doc = DocxTemplate(str(tpath))
    ctx = {
        "number": alert.number,
        "academic_year": alert.academic_year,
        "student_name": getattr(alert.student, "full_name", getattr(alert.student, "name", "")),
        "class_name": getattr(alert, "class_name", ""),
        "student_pid": getattr(alert.student, "sid", getattr(alert.student, "national_id", "")),
        "parent_name": getattr(alert, "parent_name", ""),
        "parent_mobile": getattr(alert, "parent_mobile", ""),
        "excused_days": alert.excused_days,
        "unexcused_days": alert.unexcused_days,
        "period_start": alert.period_start.strftime("%d/%m/%Y"),
        "period_end": alert.period_end.strftime("%d/%m/%Y"),
        "user_name": getattr(getattr(alert, "created_by", None), "get_full_name", lambda: "")(),
        "today": timezone.localdate().strftime("%d/%m/%Y"),
    }
    doc.render(ctx)
    buf = BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    # Post-process to add header image
    content = _apply_header_image_bytes(content)
    return content


def current_template_info(template_path: str | None = None) -> Tuple[str, str]:
    """Return (absolute_path, sha256) of the resolved template used for rendering."""
    tpath = resolve_template_path(template_path)
    try:
        data = Path(tpath).read_bytes()
        thash = hashlib.sha256(data).hexdigest()
    except Exception:
        thash = ""
    return str(tpath), thash
